# app/main.py 完整版本（关联学生学号）
from fastapi import FastAPI, Depends, Body, UploadFile, File, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from sqlalchemy import create_engine, text
from sqlalchemy.orm import sessionmaker, Session
from datetime import datetime, timedelta
from io import BytesIO
import jwt
import bcrypt
import json
import os
import uuid
import mimetypes
import zipfile
import requests
import dashscope
from dashscope import MultiModalConversation, Generation
from typing import List, Optional
from openpyxl import load_workbook
from pypdf import PdfReader
from docx import Document
from fastapi.security import OAuth2PasswordBearer  # 关键：导入OAuth2PasswordBearer
from xml.etree import ElementTree as ET

# ========== 数据库模型导入 & 配置 ==========
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy import Column, Integer, String, Boolean, DateTime, ForeignKey, Text, Float
from urllib import request as urllib_request
from urllib import error as urllib_error

# 基础模型类
Base = declarative_base()

# 1. 学生用户表（注册/登录）
class StudentUser(Base):
    __tablename__ = "student_users"
    id = Column(Integer, primary_key=True, index=True)
    name = Column(String(50), nullable=False, comment="学生姓名")
    student_id = Column(String(20), unique=True, nullable=False, comment="学号（唯一）")
    password = Column(String(100), nullable=False, comment="加密后的密码")
    default_password = Column(String(100), nullable=True, comment="默认密码")
    create_time = Column(DateTime, default=datetime.now, comment="注册时间")
    update_time = Column(DateTime, default=datetime.now, comment="更新时间")
    is_active = Column(Boolean, default=True, comment="是否激活")
    is_whitelisted = Column(Boolean, default=False, comment="是否在白名单")
    must_change_password = Column(Boolean, default=False, comment="是否需修改密码")

class AdminPermission(Base):
    __tablename__ = "admin_permissions"
    id = Column(Integer, primary_key=True, index=True)
    name = Column(String(100), nullable=False, comment="权限名称")
    key = Column(String(100), unique=True, nullable=False, comment="权限标识")
    description = Column(String(200), nullable=True, comment="权限描述")
    create_time = Column(DateTime, default=datetime.now, comment="创建时间")

class AdminRole(Base):
    __tablename__ = "admin_roles"
    id = Column(Integer, primary_key=True, index=True)
    name = Column(String(100), unique=True, nullable=False, comment="角色名称")
    description = Column(String(200), nullable=True, comment="角色描述")
    permissions = Column(Text, nullable=True, comment="权限列表JSON")
    create_time = Column(DateTime, default=datetime.now, comment="创建时间")

class AdminUser(Base):
    __tablename__ = "admin_users"
    id = Column(Integer, primary_key=True, index=True)
    username = Column(String(100), unique=True, nullable=False, comment="用户名")
    name = Column(String(100), nullable=False, comment="姓名")
    email = Column(String(200), nullable=True, comment="邮箱")
    password = Column(String(100), nullable=False, comment="加密后的密码")
    role_id = Column(Integer, ForeignKey("admin_roles.id"), nullable=True, comment="角色ID")
    is_active = Column(Boolean, default=True, comment="是否启用")
    create_time = Column(DateTime, default=datetime.now, comment="创建时间")

# 2. 成果主表（关联学生学号）
class StudentAchievement(Base):
    __tablename__ = "student_achievements"
    id = Column(Integer, primary_key=True, index=True)
    student_id = Column(String(20), nullable=False, comment="关联的学生学号")
    openid = Column(String(100), nullable=True, comment="小程序OpenID（可选）")
    create_time = Column(DateTime, default=datetime.now, comment="提交时间")
    audit_status = Column(Boolean, default=False, comment="审核状态：False未审核/True已审核")
    audit_note = Column(String(200), nullable=True, comment="审核备注")
    audit_time = Column(DateTime, nullable=True, comment="审核时间")
    overall_score = Column(Float, nullable=True, comment="总分")
    review_completed = Column(Boolean, default=False, comment="是否完成逐项审核")

# 3. 论文表
class Paper(Base):
    __tablename__ = "paper"
    id = Column(Integer, primary_key=True, index=True)
    achievement_id = Column(Integer, ForeignKey("student_achievements.id"), nullable=False, comment="关联成果主表ID")
    self_score = Column(Float, nullable=True, comment="自评分")
    title = Column(String(200), nullable=False, comment="论文标题")
    journal = Column(String(100), nullable=True, comment="发表期刊")
    publish_date = Column(String(20), nullable=True, comment="发表日期")
    review_score = Column(Float, nullable=True, comment="评分")
    rescore_score = Column(Float, nullable=True, comment="复核评分")
    review_status = Column(String(20), default="pending", comment="审核状态")
    review_comment = Column(String(500), nullable=True, comment="审核意见")
    review_time = Column(DateTime, nullable=True, comment="审核时间")
    student_agree = Column(Boolean, nullable=True, comment="学生是否同意")
    student_feedback_comment = Column(String(500), nullable=True, comment="学生反馈")
    feedback_time = Column(DateTime, nullable=True, comment="反馈时间")
    rescore_comment = Column(String(500), nullable=True, comment="复核说明")

# 4. 资政报告表
class PolicyReport(Base):
    __tablename__ = "policy_report"
    id = Column(Integer, primary_key=True, index=True)
    achievement_id = Column(Integer, ForeignKey("student_achievements.id"), nullable=False, comment="关联成果主表ID")
    self_score = Column(Float, nullable=True, comment="自评分")
    title = Column(String(200), nullable=False, comment="报告标题")
    adopt_unit = Column(String(100), nullable=True, comment="采纳单位")
    submit_date = Column(String(20), nullable=True, comment="提交日期")
    review_score = Column(Float, nullable=True, comment="评分")
    rescore_score = Column(Float, nullable=True, comment="复核评分")
    review_status = Column(String(20), default="pending", comment="审核状态")
    review_comment = Column(String(500), nullable=True, comment="审核意见")
    review_time = Column(DateTime, nullable=True, comment="审核时间")
    student_agree = Column(Boolean, nullable=True, comment="学生是否同意")
    student_feedback_comment = Column(String(500), nullable=True, comment="学生反馈")
    feedback_time = Column(DateTime, nullable=True, comment="反馈时间")
    rescore_comment = Column(String(500), nullable=True, comment="复核说明")

# 5. 学术交流表
class AcademicExchange(Base):
    __tablename__ = "academic_exchange"
    id = Column(Integer, primary_key=True, index=True)
    achievement_id = Column(Integer, ForeignKey("student_achievements.id"), nullable=False, comment="关联成果主表ID")
    self_score = Column(Float, nullable=True, comment="自评分")
    name = Column(String(200), nullable=False, comment="交流名称")
    participate_type = Column(String(50), nullable=True, comment="参与类型")
    exchange_date = Column(String(20), nullable=True, comment="交流日期")
    review_score = Column(Float, nullable=True, comment="评分")
    rescore_score = Column(Float, nullable=True, comment="复核评分")
    review_status = Column(String(20), default="pending", comment="审核状态")
    review_comment = Column(String(500), nullable=True, comment="审核意见")
    review_time = Column(DateTime, nullable=True, comment="审核时间")
    student_agree = Column(Boolean, nullable=True, comment="学生是否同意")
    student_feedback_comment = Column(String(500), nullable=True, comment="学生反馈")
    feedback_time = Column(DateTime, nullable=True, comment="反馈时间")
    rescore_comment = Column(String(500), nullable=True, comment="复核说明")

# 6. 志愿服务表
class VolunteerService(Base):
    __tablename__ = "volunteer_service"
    id = Column(Integer, primary_key=True, index=True)
    achievement_id = Column(Integer, ForeignKey("student_achievements.id"), nullable=False, comment="关联成果主表ID")
    self_score = Column(Float, nullable=True, comment="自评分")
    project_name = Column(String(200), nullable=False, comment="项目名称")
    hours = Column(Integer, default=0, comment="服务时长")
    service_date = Column(String(20), nullable=True, comment="服务日期")
    review_score = Column(Float, nullable=True, comment="评分")
    rescore_score = Column(Float, nullable=True, comment="复核评分")
    review_status = Column(String(20), default="pending", comment="审核状态")
    review_comment = Column(String(500), nullable=True, comment="审核意见")
    review_time = Column(DateTime, nullable=True, comment="审核时间")
    student_agree = Column(Boolean, nullable=True, comment="学生是否同意")
    student_feedback_comment = Column(String(500), nullable=True, comment="学生反馈")
    feedback_time = Column(DateTime, nullable=True, comment="反馈时间")
    rescore_comment = Column(String(500), nullable=True, comment="复核说明")

# 7. 获奖表
class Award(Base):
    __tablename__ = "award"
    id = Column(Integer, primary_key=True, index=True)
    achievement_id = Column(Integer, ForeignKey("student_achievements.id"), nullable=False, comment="关联成果主表ID")
    self_score = Column(Float, nullable=True, comment="自评分")
    name = Column(String(200), nullable=False, comment="奖项名称")
    level = Column(String(50), nullable=True, comment="奖项级别")
    award_date = Column(String(20), nullable=True, comment="获奖日期")
    review_score = Column(Float, nullable=True, comment="评分")
    rescore_score = Column(Float, nullable=True, comment="复核评分")
    review_status = Column(String(20), default="pending", comment="审核状态")
    review_comment = Column(String(500), nullable=True, comment="审核意见")
    review_time = Column(DateTime, nullable=True, comment="审核时间")
    student_agree = Column(Boolean, nullable=True, comment="学生是否同意")
    student_feedback_comment = Column(String(500), nullable=True, comment="学生反馈")
    feedback_time = Column(DateTime, nullable=True, comment="反馈时间")
    rescore_comment = Column(String(500), nullable=True, comment="复核说明")

# 8. 图片表
class AchievementImage(Base):
    __tablename__ = "achievement_image"
    id = Column(Integer, primary_key=True, index=True)
    achievement_id = Column(Integer, nullable=True, comment="关联成果主表ID")
    paper_id = Column(Integer, nullable=True, comment="关联论文ID")
    policy_id = Column(Integer, nullable=True, comment="关联资政报告ID")
    academic_id = Column(Integer, nullable=True, comment="关联学术交流ID")
    volunteer_id = Column(Integer, nullable=True, comment="关联志愿服务ID")
    award_id = Column(Integer, nullable=True, comment="关联获奖ID")
    file_path = Column(String(200), nullable=False, comment="图片文件路径")

class AchievementDocument(Base):
    __tablename__ = "achievement_document"
    id = Column(Integer, primary_key=True, index=True)
    achievement_id = Column(Integer, nullable=True, comment="关联成果主表ID")
    paper_id = Column(Integer, nullable=True, comment="关联论文ID")
    policy_id = Column(Integer, nullable=True, comment="关联资政报告ID")
    academic_id = Column(Integer, nullable=True, comment="关联学术交流ID")
    volunteer_id = Column(Integer, nullable=True, comment="关联志愿服务ID")
    award_id = Column(Integer, nullable=True, comment="关联获奖ID")
    custom_id = Column(Integer, nullable=True, comment="关联自定义成果ID")
    file_path = Column(String(500), nullable=False, comment="文件路径")
    file_name = Column(String(255), nullable=True, comment="原始文件名")
    file_ext = Column(String(20), nullable=True, comment="文件后缀")
    mime_type = Column(String(100), nullable=True, comment="MIME类型")
    create_time = Column(DateTime, default=datetime.now, comment="上传时间")

class ScoreFormula(Base):
    __tablename__ = "score_formula"
    id = Column(Integer, primary_key=True, index=True)
    weights_json = Column(Text, nullable=False, comment="权重配置JSON")
    update_time = Column(DateTime, default=datetime.now, comment="更新时间")

class AchievementType(Base):
    __tablename__ = "achievement_type"
    id = Column(Integer, primary_key=True, index=True)
    name = Column(String(100), nullable=False, unique=True, comment="成果类型名称")
    fields_json = Column(Text, nullable=False, comment="字段定义JSON")
    is_active = Column(Boolean, default=True, comment="是否启用")
    create_time = Column(DateTime, default=datetime.now, comment="创建时间")
    update_time = Column(DateTime, default=datetime.now, comment="更新时间")

class CustomAchievement(Base):
    __tablename__ = "custom_achievement"
    id = Column(Integer, primary_key=True, index=True)
    achievement_id = Column(Integer, ForeignKey("student_achievements.id"), nullable=False, comment="关联成果主表ID")
    self_score = Column(Float, nullable=True, comment="自评分")
    type_id = Column(Integer, ForeignKey("achievement_type.id"), nullable=False, comment="成果类型ID")
    content_json = Column(Text, nullable=False, comment="字段内容JSON")
    review_score = Column(Float, nullable=True, comment="评分")
    rescore_score = Column(Float, nullable=True, comment="复核评分")
    review_status = Column(String(20), default="pending", comment="审核状态")
    review_comment = Column(String(500), nullable=True, comment="审核意见")
    review_time = Column(DateTime, nullable=True, comment="审核时间")
    student_agree = Column(Boolean, nullable=True, comment="学生是否同意")
    student_feedback_comment = Column(String(500), nullable=True, comment="学生反馈")
    feedback_time = Column(DateTime, nullable=True, comment="反馈时间")
    rescore_comment = Column(String(500), nullable=True, comment="复核说明")

# ========== FastAPI 初始化 ==========
app = FastAPI(title="学生成果管理系统", version="1.0")

# 跨域配置
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 允许所有来源（生产环境可指定具体域名）
    allow_credentials=True,
    allow_methods=["*"],  # 允许所有HTTP方法
    allow_headers=["*"],  # 允许所有请求头
)

# ========== 数据库连接配置 ==========
# SQLite数据库路径
DATABASE_URL = "sqlite:///./student_status.db"
# 创建数据库引擎（SQLite需添加check_same_thread=False）
engine = create_engine(
    DATABASE_URL, connect_args={"check_same_thread": False}
)
# 创建会话工厂
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

# 依赖：获取数据库会话
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

def parse_permissions(raw_value: Optional[str]) -> List[str]:
    if not raw_value:
        return []
    try:
        data = json.loads(raw_value)
    except json.JSONDecodeError:
        return []
    if isinstance(data, list):
        return [str(item) for item in data]
    return []

def serialize_permission(permission: AdminPermission) -> dict:
    return {
        "id": permission.id,
        "name": permission.name,
        "key": permission.key,
        "description": permission.description or "",
        "createdAt": permission.create_time.strftime("%Y-%m-%d %H:%M:%S") if permission.create_time else ""
    }

def serialize_role(role: AdminRole) -> dict:
    return {
        "id": role.id,
        "name": role.name,
        "description": role.description or "",
        "permissions": parse_permissions(role.permissions),
        "createdAt": role.create_time.strftime("%Y-%m-%d %H:%M:%S") if role.create_time else ""
    }

def serialize_admin_user(user: AdminUser, role: Optional[AdminRole]) -> dict:
    return {
        "id": user.id,
        "username": user.username,
        "name": user.name,
        "email": user.email or "",
        "role_id": role.id if role else None,
        "role_name": role.name if role else "",
        "permissions": parse_permissions(role.permissions) if role else [],
        "createdAt": user.create_time.strftime("%Y-%m-%d %H:%M:%S") if user.create_time else "",
        "status": bool(user.is_active)
    }

def serialize_whitelist_student(student: StudentUser) -> dict:
    return {
        "student_id": student.student_id,
        "name": student.name,
        "is_active": bool(student.is_active),
        "is_whitelisted": bool(student.is_whitelisted),
        "must_change_password": bool(student.must_change_password),
        "default_password": student.default_password or "",
        "create_time": student.create_time.strftime("%Y-%m-%d %H:%M:%S") if student.create_time else "",
        "update_time": student.update_time.strftime("%Y-%m-%d %H:%M:%S") if student.update_time else ""
    }

def parse_active_value(raw_value) -> bool:
    if raw_value is None:
        return True
    value = str(raw_value).strip().lower()
    if value in ["", "1", "true", "yes", "y", "启用", "是"]:
        return True
    if value in ["0", "false", "no", "n", "禁用", "否"]:
        return False
    return True

def upsert_whitelist_student(
    db: Session,
    student_id: str,
    name: str,
    default_password: str,
    is_active: bool = True
):
    student = db.query(StudentUser).filter(StudentUser.student_id == student_id).first()
    password_bytes = default_password.encode("utf-8")
    hashed_password = bcrypt.hashpw(password_bytes, bcrypt.gensalt()).decode("utf-8")
    created = False
    if student:
        student.name = name
        student.default_password = default_password
        student.password = hashed_password
        student.is_whitelisted = True
        student.is_active = is_active
        student.must_change_password = True
        student.update_time = datetime.now()
    else:
        student = StudentUser(
            name=name,
            student_id=student_id,
            password=hashed_password,
            default_password=default_password,
            create_time=datetime.now(),
            update_time=datetime.now(),
            is_active=is_active,
            is_whitelisted=True,
            must_change_password=True
        )
        db.add(student)
        created = True
    return student, created

REVIEW_COLUMNS = [
    ("review_score", "FLOAT"),
    ("rescore_score", "FLOAT"),
    ("review_status", "VARCHAR(20) DEFAULT 'pending'"),
    ("review_comment", "VARCHAR(500)"),
    ("review_time", "DATETIME"),
    ("student_agree", "BOOLEAN"),
    ("student_feedback_comment", "VARCHAR(500)"),
    ("feedback_time", "DATETIME"),
    ("rescore_comment", "VARCHAR(500)")
]

REVIEW_MODEL_MAP = {
    "paper": (Paper, "title"),
    "policy": (PolicyReport, "title"),
    "academic": (AcademicExchange, "name"),
    "volunteer": (VolunteerService, "project_name"),
    "award": (Award, "name"),
    "custom": (CustomAchievement, "content_json")
}

def get_or_init_score_formula(db: Session) -> ScoreFormula:
    formula = db.query(ScoreFormula).order_by(ScoreFormula.id.asc()).first()
    if formula:
        return formula
    default_weights = {
        "paper": 0.2,
        "policy": 0.2,
        "academic": 0.2,
        "volunteer": 0.2,
        "award": 0.2,
        "custom": 0.0
    }
    formula = ScoreFormula(
        weights_json=json.dumps(default_weights, ensure_ascii=False),
        update_time=datetime.now()
    )
    db.add(formula)
    db.commit()
    db.refresh(formula)
    return formula

def parse_weights(raw_value: Optional[str]) -> dict:
    if not raw_value:
        return {}
    try:
        data = json.loads(raw_value)
    except Exception:
        return {}
    if isinstance(data, dict):
        return data
    return {}

def ensure_student_user_schema():
    with engine.begin() as conn:
        table_rows = conn.execute(text("PRAGMA table_info(student_users)")).fetchall()
        existing_columns = {row[1] for row in table_rows}
        if "default_password" not in existing_columns:
            conn.execute(text("ALTER TABLE student_users ADD COLUMN default_password VARCHAR(100)"))
        if "update_time" not in existing_columns:
            conn.execute(text("ALTER TABLE student_users ADD COLUMN update_time DATETIME"))
        if "is_whitelisted" not in existing_columns:
            conn.execute(text("ALTER TABLE student_users ADD COLUMN is_whitelisted BOOLEAN DEFAULT 0"))
        if "must_change_password" not in existing_columns:
            conn.execute(text("ALTER TABLE student_users ADD COLUMN must_change_password BOOLEAN DEFAULT 0"))

def ensure_table_columns(conn, table_name: str, columns):
    table_rows = conn.execute(text(f"PRAGMA table_info({table_name})")).fetchall()
    existing_columns = {row[1] for row in table_rows}
    for column_name, column_sql in columns:
        if column_name not in existing_columns:
            conn.execute(text(f"ALTER TABLE {table_name} ADD COLUMN {column_name} {column_sql}"))

def ensure_achievement_schema():
    with engine.begin() as conn:
        ensure_table_columns(conn, "student_achievements", [
            ("overall_score", "FLOAT"),
            ("review_completed", "BOOLEAN DEFAULT 0")
        ])
        for table_name in ["paper", "policy_report", "academic_exchange", "volunteer_service", "award", "custom_achievement"]:
            ensure_table_columns(conn, table_name, REVIEW_COLUMNS)

def serialize_document(doc: AchievementDocument) -> dict:
    return {
        "id": doc.id,
        "file_path": doc.file_path,
        "file_name": doc.file_name or "",
        "file_ext": doc.file_ext or "",
        "mime_type": doc.mime_type or "",
        "download_url": f"/uploads/{doc.file_path}"
    }

def get_documents_for_item(db: Session, item_type: str, item_id: int) -> List[dict]:
    query = db.query(AchievementDocument)
    if item_type == "paper":
        query = query.filter(AchievementDocument.paper_id == item_id)
    elif item_type == "policy":
        query = query.filter(AchievementDocument.policy_id == item_id)
    elif item_type == "academic":
        query = query.filter(AchievementDocument.academic_id == item_id)
    elif item_type == "volunteer":
        query = query.filter(AchievementDocument.volunteer_id == item_id)
    elif item_type == "award":
        query = query.filter(AchievementDocument.award_id == item_id)
    elif item_type == "custom":
        query = query.filter(AchievementDocument.custom_id == item_id)
    else:
        return []
    return [serialize_document(item) for item in query.order_by(AchievementDocument.id.asc()).all()]

def append_documents_for_item(db: Session, achievement_id: int, item_type: str, item_id: int, docs: list):
    for entry in docs or []:
        if isinstance(entry, str):
            file_path = entry
            file_name = os.path.basename(entry)
        elif isinstance(entry, dict):
            file_path = str(entry.get("file_path") or entry.get("path") or "").strip()
            file_name = str(entry.get("file_name") or os.path.basename(file_path)).strip()
        else:
            continue
        if not file_path:
            continue
        file_ext = os.path.splitext(file_name or file_path)[1].replace(".", "").lower()
        mime_type = mimetypes.guess_type(file_name or file_path)[0] or ""
        payload = {
            "achievement_id": achievement_id,
            "file_path": file_path,
            "file_name": file_name or os.path.basename(file_path),
            "file_ext": file_ext,
            "mime_type": mime_type
        }
        if item_type == "paper":
            payload["paper_id"] = item_id
        elif item_type == "policy":
            payload["policy_id"] = item_id
        elif item_type == "academic":
            payload["academic_id"] = item_id
        elif item_type == "volunteer":
            payload["volunteer_id"] = item_id
        elif item_type == "award":
            payload["award_id"] = item_id
        elif item_type == "custom":
            payload["custom_id"] = item_id
        db.add(AchievementDocument(**payload))

def calculate_review_completed(achievement_data: dict) -> bool:
    for key in ["papers", "policies", "academics", "volunteers", "awards", "customs"]:
        for item in achievement_data.get(key, []):
            review_status = item.get("review_status")
            review_score = item.get("rescore_score") if item.get("review_status") == "rescored" else item.get("review_score")
            if review_score is None and item.get("review_score") is None:
                return False
            if review_status not in ["reviewed", "rescored", "agreed"]:
                return False
    return True

def get_type_items(db: Session, achievement_id: int, item_type: str):
    item_type = (item_type or "").lower()
    model_info = REVIEW_MODEL_MAP.get(item_type)
    if not model_info:
        return []
    model_cls = model_info[0]
    return db.query(model_cls).filter(model_cls.achievement_id == achievement_id).all()

def build_type_summary(display_name: str, item_type: str, items: list):
    if not items:
        return None
    first_item = items[0]
    review_score = first_item.review_score
    rescore_score = first_item.rescore_score
    score = rescore_score if rescore_score is not None else review_score
    review_comment = first_item.review_comment or ""
    rescore_comment = first_item.rescore_comment or ""
    any_disagree = any(item.student_agree is False for item in items)
    pending_rescore_items = [item for item in items if item.student_agree is False]
    all_disagree_rescored = bool(pending_rescore_items) and all(
        (item.review_status or "") == "rescored" and bool((item.rescore_comment or "").strip())
        for item in pending_rescore_items
    )
    has_reviewed = any(
        item.review_score is not None or (item.review_status or "") in ["reviewed", "agreed", "rescored"]
        for item in items
    )
    if all_disagree_rescored:
        indicator_color = "yellow"
        indicator_text = "已复核"
        audit_status = "已复核"
    elif any_disagree:
        indicator_color = "red"
        indicator_text = "待复核"
        audit_status = "待复核"
    elif has_reviewed:
        indicator_color = "primary"
        indicator_text = "已审核"
        audit_status = "已审核"
    else:
        indicator_color = "info"
        indicator_text = "已提交"
        audit_status = "已提交"
    return {
        "id": first_item.id,
        "type": item_type,
        "type_name": display_name,
        "count": len(items),
        "score": score,
        "review_score": review_score,
        "rescore_score": rescore_score,
        "review_status": first_item.review_status or "pending",
        "review_comment": review_comment,
        "rescore_comment": rescore_comment,
        "student_agree": False if any_disagree else None,
        "student_feedback_comment": first_item.student_feedback_comment or "",
        "audit_status": audit_status,
        "indicator_color": indicator_color,
        "indicator_text": indicator_text
    }

def calculate_achievement_lifecycle_status(db: Session, achievement: StudentAchievement) -> str:
    source_items = [
        db.query(Paper).filter(Paper.achievement_id == achievement.id).all(),
        db.query(PolicyReport).filter(PolicyReport.achievement_id == achievement.id).all(),
        db.query(AcademicExchange).filter(AcademicExchange.achievement_id == achievement.id).all(),
        db.query(VolunteerService).filter(VolunteerService.achievement_id == achievement.id).all(),
        db.query(Award).filter(Award.achievement_id == achievement.id).all(),
        db.query(CustomAchievement).filter(CustomAchievement.achievement_id == achievement.id).all()
    ]
    items = [item for group in source_items for item in group]
    if not items:
        return "已提交"
    disagree_items = [item for item in items if item.student_agree is False]
    if disagree_items:
        all_rescored = all(
            (item.review_status or "") == "rescored" and bool((item.rescore_comment or "").strip())
            for item in disagree_items
        )
        return "已复核" if all_rescored else "待复核"
    has_reviewed = any(
        item.review_score is not None or (item.review_status or "") in ["reviewed", "agreed", "rescored"]
        for item in items
    )
    return "已审核" if has_reviewed else "已提交"

def recalculate_achievement_score(db: Session, achievement: StudentAchievement):
    formula = get_or_init_score_formula(db)
    weights = parse_weights(formula.weights_json)
    source_items = {
        "paper": db.query(Paper).filter(Paper.achievement_id == achievement.id).all(),
        "policy": db.query(PolicyReport).filter(PolicyReport.achievement_id == achievement.id).all(),
        "academic": db.query(AcademicExchange).filter(AcademicExchange.achievement_id == achievement.id).all(),
        "volunteer": db.query(VolunteerService).filter(VolunteerService.achievement_id == achievement.id).all(),
        "award": db.query(Award).filter(Award.achievement_id == achievement.id).all(),
        "custom": db.query(CustomAchievement).filter(CustomAchievement.achievement_id == achievement.id).all()
    }
    weighted_total = 0.0
    weighted_factor = 0.0
    all_reviewed = True
    for key, items in source_items.items():
        if not items:
            continue
        for item in items:
            if (item.review_status or "pending") not in ["reviewed", "rescored", "agreed"]:
                all_reviewed = False
        scores = []
        for item in items:
            effective_score = item.rescore_score if item.rescore_score is not None else item.review_score
            if effective_score is None:
                continue
            scores.append(float(effective_score))
        if len(scores) != len(items):
            all_reviewed = False
        if not scores:
            continue
        weight = float(weights.get(key, 0))
        if weight <= 0:
            continue
        avg_score = sum(scores) / len(scores)
        weighted_total += avg_score * weight
        weighted_factor += weight
    achievement.overall_score = (weighted_total / weighted_factor) if weighted_factor > 0 else None
    achievement.review_completed = bool(all_reviewed and weighted_factor > 0)

def parse_suggestion_json(content) -> dict:
    def normalize(data: dict) -> dict:
        if not isinstance(data, dict):
            return {}
        if isinstance(data.get("suggestions"), dict):
            return data.get("suggestions")
        if isinstance(data.get("data"), dict) and isinstance(data["data"].get("suggestions"), dict):
            return data["data"]["suggestions"]
        return data
    if isinstance(content, dict):
        return normalize(content)
    if not isinstance(content, str):
        return {}
    text_content = content.strip()
    if not text_content:
        return {}
    try:
        data = json.loads(text_content)
        if isinstance(data, dict):
            return normalize(data)
    except Exception:
        pass
    if "```" in text_content:
        blocks = text_content.split("```")
        for block in blocks:
            candidate = block.strip()
            if candidate.lower().startswith("json"):
                candidate = candidate[4:].strip()
            try:
                data = json.loads(candidate)
                if isinstance(data, dict):
                    return normalize(data)
            except Exception:
                continue
    left = text_content.find("{")
    right = text_content.rfind("}")
    if left >= 0 and right > left:
        snippet = text_content[left:right + 1]
        try:
            data = json.loads(snippet)
            if isinstance(data, dict):
                return normalize(data)
        except Exception:
            return {}
    return {}

# ========== JWT 配置（登录Token） ==========
SECRET_KEY = "your-secret-key-20260221"  # 替换为随机字符串（建议用：openssl rand -hex 32）
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 120  # Token有效期2小时
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/student/login")
# ========== 图片上传配置 ==========
# 上传目录（确保存在）
UPLOAD_DIR = "./uploads"
if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)

# ========== 接口定义 ==========
# 1. 测试接口
@app.get("/test")
async def test_api():
    return {"message": "服务器正常运行", "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

# 2. 学生注册接口
@app.post("/student/register")
async def student_register(
    name: str = Body(...),
    student_id: str = Body(...),
    password: str = Body(...),
    db: Session = Depends(get_db)
):
    try:
        existing_student = db.query(StudentUser).filter(StudentUser.student_id == student_id).first()
        if not existing_student or not existing_student.is_whitelisted:
            return {
                "success": False,
                "message": "该学号不在白名单，请联系管理员"
            }
        password_bytes = password.encode('utf-8')
        salt = bcrypt.gensalt()
        hashed_password = bcrypt.hashpw(password_bytes, salt).decode('utf-8')
        existing_student.name = name or existing_student.name
        existing_student.password = hashed_password
        existing_student.is_active = True
        existing_student.must_change_password = False
        existing_student.update_time = datetime.now()
        db.commit()
        return {
            "success": True,
            "message": "账号开通成功，请登录",
            "student_id": student_id
        }
    except Exception as e:
        db.rollback()
        print(f"注册失败：{str(e)}")
        return {
            "success": False,
            "message": f"注册失败：{str(e)}"
        }

# 3. 学生登录接口
@app.post("/student/login")
async def student_login(
    student_id: str = Body(...),
    password: str = Body(...),
    db: Session = Depends(get_db)
):
    try:
        student = db.query(StudentUser).filter(
            StudentUser.student_id == student_id,
            StudentUser.is_whitelisted == True,
            StudentUser.is_active == True
        ).first()
        if not student:
            return {
                "success": False,
                "message": "账号不存在、未启用或不在白名单"
            }
        password_bytes = password.encode('utf-8')
        hashed_password_bytes = student.password.encode('utf-8')
        if not bcrypt.checkpw(password_bytes, hashed_password_bytes):
            return {
                "success": False,
                "message": "密码错误"
            }
        access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        access_token = jwt.encode(
            {
                "sub": student.student_id,
                "name": student.name,
                "exp": datetime.utcnow() + access_token_expires
            },
            SECRET_KEY,
            algorithm=ALGORITHM
        )
        
        return {
            "success": True,
            "message": "登录成功",
            "token": access_token,
            "student_id": student.student_id,
            "name": student.name,
            "must_change_password": bool(student.must_change_password)
        }
    except Exception as e:
        print(f"登录失败：{str(e)}")
        return {
            "success": False,
            "message": f"登录失败：{str(e)}"
        }

@app.post("/admin/login")
async def admin_login(
    username: str = Body(...),
    password: str = Body(...),
    db: Session = Depends(get_db)
):
    try:
        user = db.query(AdminUser).filter(AdminUser.username == username).first()
        if not user or not user.is_active:
            return {
                "success": False,
                "message": "用户不存在或已禁用"
            }
        password_bytes = password.encode("utf-8")
        hashed_password_bytes = user.password.encode("utf-8")
        if not bcrypt.checkpw(password_bytes, hashed_password_bytes):
            return {
                "success": False,
                "message": "密码错误"
            }
        role = db.query(AdminRole).filter(AdminRole.id == user.role_id).first()
        access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        access_token = jwt.encode(
            {
                "sub": user.username,
                "role_id": role.id if role else None,
                "exp": datetime.utcnow() + access_token_expires
            },
            SECRET_KEY,
            algorithm=ALGORITHM
        )
        return {
            "success": True,
            "message": "登录成功",
            "token": access_token,
            "user": serialize_admin_user(user, role)
        }
    except Exception as e:
        return {
            "success": False,
            "message": f"登录失败：{str(e)}"
        }

@app.get("/admin/permissions")
async def get_permissions(db: Session = Depends(get_db)):
    permissions = db.query(AdminPermission).order_by(AdminPermission.id.asc()).all()
    return {
        "code": 200,
        "data": {
            "list": [serialize_permission(item) for item in permissions]
        },
        "message": "查询成功"
    }

@app.post("/admin/permissions")
async def create_permission(data: dict = Body(...), db: Session = Depends(get_db)):
    name = (data.get("name") or "").strip()
    key = (data.get("key") or "").strip()
    if not name or not key:
        return {
            "code": 400,
            "message": "权限名称和标识不能为空"
        }
    exists = db.query(AdminPermission).filter(AdminPermission.key == key).first()
    if exists:
        return {
            "code": 400,
            "message": "权限标识已存在"
        }
    permission = AdminPermission(
        name=name,
        key=key,
        description=data.get("description", ""),
        create_time=datetime.now()
    )
    db.add(permission)
    db.commit()
    db.refresh(permission)
    return {
        "code": 200,
        "data": serialize_permission(permission),
        "message": "新增成功"
    }

@app.put("/admin/permissions/{permission_id}")
async def update_permission(
    permission_id: int,
    data: dict = Body(...),
    db: Session = Depends(get_db)
):
    permission = db.query(AdminPermission).filter(AdminPermission.id == permission_id).first()
    if not permission:
        return {
            "code": 404,
            "message": "权限不存在"
        }
    name = (data.get("name") or "").strip()
    key = (data.get("key") or "").strip()
    if not name or not key:
        return {
            "code": 400,
            "message": "权限名称和标识不能为空"
        }
    exists = db.query(AdminPermission).filter(
        AdminPermission.key == key,
        AdminPermission.id != permission_id
    ).first()
    if exists:
        return {
            "code": 400,
            "message": "权限标识已存在"
        }
    permission.name = name
    permission.key = key
    permission.description = data.get("description", "")
    db.commit()
    db.refresh(permission)
    return {
        "code": 200,
        "data": serialize_permission(permission),
        "message": "更新成功"
    }

@app.delete("/admin/permissions/{permission_id}")
async def delete_permission(permission_id: int, db: Session = Depends(get_db)):
    permission = db.query(AdminPermission).filter(AdminPermission.id == permission_id).first()
    if not permission:
        return {
            "code": 404,
            "message": "权限不存在"
        }
    db.delete(permission)
    db.commit()
    return {
        "code": 200,
        "message": "删除成功"
    }

@app.get("/admin/roles")
async def get_roles(db: Session = Depends(get_db)):
    roles = db.query(AdminRole).order_by(AdminRole.id.asc()).all()
    return {
        "code": 200,
        "data": {
            "list": [serialize_role(item) for item in roles]
        },
        "message": "查询成功"
    }

@app.post("/admin/roles")
async def create_role(data: dict = Body(...), db: Session = Depends(get_db)):
    name = (data.get("name") or "").strip()
    if not name:
        return {
            "code": 400,
            "message": "角色名称不能为空"
        }
    exists = db.query(AdminRole).filter(AdminRole.name == name).first()
    if exists:
        return {
            "code": 400,
            "message": "角色名称已存在"
        }
    permissions = data.get("permissions", [])
    role = AdminRole(
        name=name,
        description=data.get("description", ""),
        permissions=json.dumps(permissions, ensure_ascii=False),
        create_time=datetime.now()
    )
    db.add(role)
    db.commit()
    db.refresh(role)
    return {
        "code": 200,
        "data": serialize_role(role),
        "message": "新增成功"
    }

@app.put("/admin/roles/{role_id}")
async def update_role(role_id: int, data: dict = Body(...), db: Session = Depends(get_db)):
    role = db.query(AdminRole).filter(AdminRole.id == role_id).first()
    if not role:
        return {
            "code": 404,
            "message": "角色不存在"
        }
    name = (data.get("name") or "").strip()
    if not name:
        return {
            "code": 400,
            "message": "角色名称不能为空"
        }
    exists = db.query(AdminRole).filter(
        AdminRole.name == name,
        AdminRole.id != role_id
    ).first()
    if exists:
        return {
            "code": 400,
            "message": "角色名称已存在"
        }
    role.name = name
    role.description = data.get("description", "")
    if "permissions" in data:
        role.permissions = json.dumps(data.get("permissions") or [], ensure_ascii=False)
    db.commit()
    db.refresh(role)
    return {
        "code": 200,
        "data": serialize_role(role),
        "message": "更新成功"
    }

@app.delete("/admin/roles/{role_id}")
async def delete_role(role_id: int, db: Session = Depends(get_db)):
    role = db.query(AdminRole).filter(AdminRole.id == role_id).first()
    if not role:
        return {
            "code": 404,
            "message": "角色不存在"
        }
    db.delete(role)
    db.commit()
    return {
        "code": 200,
        "message": "删除成功"
    }

@app.get("/admin/users")
async def get_admin_users(db: Session = Depends(get_db)):
    users = db.query(AdminUser).order_by(AdminUser.id.asc()).all()
    role_map = {role.id: role for role in db.query(AdminRole).all()}
    return {
        "code": 200,
        "data": {
            "list": [serialize_admin_user(user, role_map.get(user.role_id)) for user in users]
        },
        "message": "查询成功"
    }

@app.post("/admin/users")
async def create_admin_user(data: dict = Body(...), db: Session = Depends(get_db)):
    username = (data.get("username") or "").strip()
    name = (data.get("name") or "").strip()
    password = (data.get("password") or "").strip()
    role_id = data.get("role_id")
    if not username or not name or not password:
        return {
            "code": 400,
            "message": "用户名、姓名、密码不能为空"
        }
    exists = db.query(AdminUser).filter(AdminUser.username == username).first()
    if exists:
        return {
            "code": 400,
            "message": "用户名已存在"
        }
    role = db.query(AdminRole).filter(AdminRole.id == role_id).first() if role_id else None
    if role_id and not role:
        return {
            "code": 400,
            "message": "角色不存在"
        }
    password_bytes = password.encode("utf-8")
    hashed_password = bcrypt.hashpw(password_bytes, bcrypt.gensalt()).decode("utf-8")
    user = AdminUser(
        username=username,
        name=name,
        email=data.get("email", ""),
        password=hashed_password,
        role_id=role.id if role else None,
        is_active=True,
        create_time=datetime.now()
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    return {
        "code": 200,
        "data": serialize_admin_user(user, role),
        "message": "新增成功"
    }

@app.put("/admin/users/{user_id}")
async def update_admin_user(
    user_id: int,
    data: dict = Body(...),
    db: Session = Depends(get_db)
):
    user = db.query(AdminUser).filter(AdminUser.id == user_id).first()
    if not user:
        return {
            "code": 404,
            "message": "用户不存在"
        }
    username = (data.get("username") or "").strip()
    name = (data.get("name") or "").strip()
    if not username or not name:
        return {
            "code": 400,
            "message": "用户名和姓名不能为空"
        }
    exists = db.query(AdminUser).filter(
        AdminUser.username == username,
        AdminUser.id != user_id
    ).first()
    if exists:
        return {
            "code": 400,
            "message": "用户名已存在"
        }
    role_id = data.get("role_id")
    role = db.query(AdminRole).filter(AdminRole.id == role_id).first() if role_id else None
    if role_id and not role:
        return {
            "code": 400,
            "message": "角色不存在"
        }
    user.username = username
    user.name = name
    user.email = data.get("email", "")
    user.role_id = role.id if role else None
    if "is_active" in data:
        user.is_active = bool(data.get("is_active"))
    new_password = (data.get("password") or "").strip()
    if new_password:
        password_bytes = new_password.encode("utf-8")
        user.password = bcrypt.hashpw(password_bytes, bcrypt.gensalt()).decode("utf-8")
    db.commit()
    db.refresh(user)
    return {
        "code": 200,
        "data": serialize_admin_user(user, role),
        "message": "更新成功"
    }

@app.delete("/admin/users/{user_id}")
async def delete_admin_user(user_id: int, db: Session = Depends(get_db)):
    user = db.query(AdminUser).filter(AdminUser.id == user_id).first()
    if not user:
        return {
            "code": 404,
            "message": "用户不存在"
        }
    db.delete(user)
    db.commit()
    return {
        "code": 200,
        "message": "删除成功"
    }

@app.get("/admin/whitelist")
async def get_whitelist_students(
    page: int = 1,
    size: int = 20,
    student_id: Optional[str] = None,
    name: Optional[str] = None,
    is_active: Optional[bool] = None,
    db: Session = Depends(get_db)
):
    query = db.query(StudentUser).filter(StudentUser.is_whitelisted == True)
    if student_id and student_id.strip():
        query = query.filter(StudentUser.student_id.like(f"%{student_id.strip()}%"))
    if name and name.strip():
        query = query.filter(StudentUser.name.like(f"%{name.strip()}%"))
    if is_active is not None:
        query = query.filter(StudentUser.is_active == is_active)
    total = query.count()
    items = query.order_by(StudentUser.create_time.desc()).offset((page - 1) * size).limit(size).all()
    return {
        "code": 200,
        "data": {
            "list": [serialize_whitelist_student(item) for item in items],
            "total": total,
            "page": page,
            "size": size
        },
        "message": "查询成功"
    }

@app.post("/admin/whitelist")
async def create_whitelist_student(data: dict = Body(...), db: Session = Depends(get_db)):
    student_id = (data.get("student_id") or "").strip()
    name = (data.get("name") or "").strip()
    default_password = (data.get("default_password") or "123456").strip()
    is_active = bool(data.get("is_active", True))
    if not student_id or not name or not default_password:
        return {
            "code": 400,
            "message": "学号、姓名、默认密码不能为空"
        }
    student, _ = upsert_whitelist_student(
        db=db,
        student_id=student_id,
        name=name,
        default_password=default_password,
        is_active=is_active
    )
    db.commit()
    db.refresh(student)
    return {
        "code": 200,
        "data": serialize_whitelist_student(student),
        "message": "保存成功"
    }

@app.post("/admin/whitelist/import")
async def import_whitelist_students(file: UploadFile = File(...), db: Session = Depends(get_db)):
    file_name = file.filename or ""
    if not file_name.lower().endswith(".xlsx"):
        return {
            "code": 400,
            "message": "仅支持 .xlsx 文件"
        }
    try:
        file_content = await file.read()
        workbook = load_workbook(filename=BytesIO(file_content), data_only=True)
        worksheet = workbook.active
        created_count = 0
        updated_count = 0
        skipped_count = 0
        errors = []
        for row_index, row in enumerate(worksheet.iter_rows(min_row=2, values_only=True), start=2):
            raw_student_id = row[0] if len(row) > 0 else None
            raw_name = row[1] if len(row) > 1 else None
            raw_default_password = row[2] if len(row) > 2 else None
            raw_is_active = row[3] if len(row) > 3 else None
            student_id = str(raw_student_id or "").strip()
            name = str(raw_name or "").strip()
            default_password = str(raw_default_password or "123456").strip()
            if not student_id and not name:
                skipped_count += 1
                continue
            if not student_id or not name:
                errors.append({"row": row_index, "message": "学号或姓名为空"})
                continue
            is_active = parse_active_value(raw_is_active)
            _, created = upsert_whitelist_student(
                db=db,
                student_id=student_id,
                name=name,
                default_password=default_password or "123456",
                is_active=is_active
            )
            if created:
                created_count += 1
            else:
                updated_count += 1
        db.commit()
        return {
            "code": 200,
            "data": {
                "created": created_count,
                "updated": updated_count,
                "skipped": skipped_count,
                "errors": errors
            },
            "message": "批量导入完成"
        }
    except Exception as e:
        db.rollback()
        return {
            "code": 500,
            "message": f"批量导入失败：{str(e)}"
        }

@app.put("/admin/whitelist/{student_id}")
async def update_whitelist_student(student_id: str, data: dict = Body(...), db: Session = Depends(get_db)):
    student = db.query(StudentUser).filter(
        StudentUser.student_id == student_id,
        StudentUser.is_whitelisted == True
    ).first()
    if not student:
        return {
            "code": 404,
            "message": "白名单学生不存在"
        }
    if "name" in data and str(data.get("name") or "").strip():
        student.name = str(data.get("name")).strip()
    if "is_active" in data:
        student.is_active = bool(data.get("is_active"))
    if "default_password" in data and str(data.get("default_password") or "").strip():
        default_password = str(data.get("default_password")).strip()
        password_bytes = default_password.encode("utf-8")
        student.default_password = default_password
        student.password = bcrypt.hashpw(password_bytes, bcrypt.gensalt()).decode("utf-8")
        student.must_change_password = True
    student.update_time = datetime.now()
    db.commit()
    db.refresh(student)
    return {
        "code": 200,
        "data": serialize_whitelist_student(student),
        "message": "更新成功"
    }

@app.put("/admin/whitelist/{student_id}/reset-password")
async def reset_whitelist_student_password(student_id: str, data: dict = Body({}), db: Session = Depends(get_db)):
    student = db.query(StudentUser).filter(
        StudentUser.student_id == student_id,
        StudentUser.is_whitelisted == True
    ).first()
    if not student:
        return {
            "code": 404,
            "message": "白名单学生不存在"
        }
    new_default_password = str(data.get("default_password") or student.default_password or "123456").strip()
    password_bytes = new_default_password.encode("utf-8")
    student.default_password = new_default_password
    student.password = bcrypt.hashpw(password_bytes, bcrypt.gensalt()).decode("utf-8")
    student.must_change_password = True
    student.update_time = datetime.now()
    db.commit()
    db.refresh(student)
    return {
        "code": 200,
        "data": serialize_whitelist_student(student),
        "message": "重置成功"
    }

@app.delete("/admin/whitelist/{student_id}")
async def remove_whitelist_student(student_id: str, db: Session = Depends(get_db)):
    student = db.query(StudentUser).filter(
        StudentUser.student_id == student_id,
        StudentUser.is_whitelisted == True
    ).first()
    if not student:
        return {
            "code": 404,
            "message": "白名单学生不存在"
        }
    student.is_whitelisted = False
    student.must_change_password = False
    student.update_time = datetime.now()
    db.commit()
    db.refresh(student)
    return {
        "code": 200,
        "data": serialize_whitelist_student(student),
        "message": "移出白名单成功"
    }

# 验证token并获取当前学生
async def get_current_student(
    token: str = Depends(oauth2_scheme),
    db: Session = Depends(get_db)
):
    credentials_exception = HTTPException(
        status_code=401,
        detail="认证失败，请重新登录",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        # 解码token
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        student_id: str = payload.get("sub")
        if student_id is None:
            raise credentials_exception
    except jwt.PyJWTError:
        raise credentials_exception
    
    # 从数据库查询学生信息
    student = db.query(StudentUser).filter(
        StudentUser.student_id == student_id,
        StudentUser.is_active == True,
        StudentUser.is_whitelisted == True
    ).first()
    
    if student is None:
        raise credentials_exception
    return student

# 获取学生信息接口（首页调用）
@app.get("/api/student/info", summary="获取当前登录学生信息")
async def get_student_info(
    current_student: StudentUser = Depends(get_current_student)
):
    # 返回前端需要的字段（对应前端的 studentName/studentId）
    return {
        "code": 200,
        "message": "获取成功",
        "data": {
            "studentName": current_student.name,    # 对应前端 studentName
            "studentId": current_student.student_id # 对应前端 studentId
        }
    }

@app.post("/student/change-password")
async def student_change_password(
    old_password: str = Body(...),
    new_password: str = Body(...),
    current_student: StudentUser = Depends(get_current_student),
    db: Session = Depends(get_db)
):
    try:
        old_password_bytes = old_password.encode("utf-8")
        hashed_password_bytes = current_student.password.encode("utf-8")
        if not bcrypt.checkpw(old_password_bytes, hashed_password_bytes):
            return {
                "success": False,
                "message": "旧密码错误"
            }
        new_password_bytes = new_password.encode("utf-8")
        current_student.password = bcrypt.hashpw(new_password_bytes, bcrypt.gensalt()).decode("utf-8")
        current_student.must_change_password = False
        current_student.update_time = datetime.now()
        db.commit()
        return {
            "success": True,
            "message": "密码修改成功"
        }
    except Exception as e:
        db.rollback()
        return {
            "success": False,
            "message": f"修改失败：{str(e)}"
        }


def save_uploaded_file(file: UploadFile):
    original_name = file.filename or "file"
    ext = os.path.splitext(original_name)[1]
    file_name = f"{uuid.uuid4()}{ext}"
    file_path = os.path.join(UPLOAD_DIR, file_name)
    content_type = file.content_type or mimetypes.guess_type(original_name)[0] or ""
    return file_name, file_path, original_name, content_type

@app.post("/upload/document")
async def upload_document(file: UploadFile = File(...)):
    try:
        file_name, file_path, original_name, content_type = save_uploaded_file(file)
        with open(file_path, "wb") as f:
            f.write(await file.read())
        return {
            "success": True,
            "file_name": original_name,
            "stored_name": file_name,
            "file_path": file_name,
            "file_ext": os.path.splitext(original_name)[1].replace(".", "").lower(),
            "mime_type": content_type,
            "message": "文件上传成功"
        }
    except Exception as e:
        print(f"文件上传失败：{str(e)}")
        return {
            "success": False,
            "message": f"文件上传失败：{str(e)}"
        }

@app.post("/upload/image")
async def upload_image(file: UploadFile = File(...)):
    return await upload_document(file)

@app.post("/submit/achievements")
async def submit_achievements(
    request: Request,
    db: Session = Depends(get_db)
):
    try:
        try:
            payload = await request.json()
        except Exception:
            payload = {}
        def parse_self_score(value):
            if value is None:
                return None
            text = str(value).strip()
            if text == "":
                return None
            try:
                return float(text)
            except Exception:
                return None

        def get_self_score_value(item):
            if not isinstance(item, dict):
                return None
            if item.get("self_score") is not None:
                return item.get("self_score")
            if item.get("selfScore") is not None:
                return item.get("selfScore")
            if item.get("selfscore") is not None:
                return item.get("selfscore")
            return None

        if not isinstance(payload, dict):
            return {"success": False, "message": "提交数据格式错误"}
        student_id = str(payload.get("student_id", "")).strip()
        paper_items = payload.get("paperList", []) or []
        policy_items = payload.get("policyList", []) or []
        academic_items = payload.get("academicList", []) or []
        volunteer_items = payload.get("volunteerList", []) or []
        award_items = payload.get("awardList", []) or []
        custom_items = payload.get("customList", []) or []
        if not student_id:
            return {"success": False, "message": "学号不能为空"}
        student = db.query(StudentUser).filter(StudentUser.student_id == student_id).first()

        if not student:
            return {
                "success": False,
                "message": "学号未注册，无法提交成果"
            }

        achievement = StudentAchievement(
            student_id=student_id,
            create_time=datetime.now(),
            audit_status=False,
            review_completed=False
        )
        db.add(achievement)
        db.commit()
        db.refresh(achievement)
        achievement_id = achievement.id

        for paper in paper_items:
            if paper.get("title"):
                paper_self_score = parse_self_score(get_self_score_value(paper))
                if paper_self_score is None:
                    return {"success": False, "message": "论文成果缺少自评分，请补充后提交"}
                new_paper = Paper(
                    achievement_id=achievement_id,
                    title=paper.get("title"),
                    journal=paper.get("journal"),
                    publish_date=paper.get("date"),
                    self_score=paper_self_score,
                    review_status="pending"
                )
                db.add(new_paper)
                db.commit()
                db.refresh(new_paper)
                all_docs = list(paper.get("documents", []) or []) + list(paper.get("images", []) or [])
                append_documents_for_item(db, achievement_id, "paper", new_paper.id, all_docs)

        for policy in policy_items:
            if policy.get("title"):
                policy_self_score = parse_self_score(get_self_score_value(policy))
                if policy_self_score is None:
                    return {"success": False, "message": "资政报告缺少自评分，请补充后提交"}
                new_policy = PolicyReport(
                    achievement_id=achievement_id,
                    title=policy.get("title"),
                    adopt_unit=policy.get("adopt_unit"),
                    submit_date=policy.get("date"),
                    self_score=policy_self_score,
                    review_status="pending"
                )
                db.add(new_policy)
                db.commit()
                db.refresh(new_policy)
                all_docs = list(policy.get("documents", []) or []) + list(policy.get("images", []) or [])
                append_documents_for_item(db, achievement_id, "policy", new_policy.id, all_docs)

        participate_types = ['参会', '报告发言', '墙报展示', '其他']
        for academic in academic_items:
            if academic.get("name"):
                academic_self_score = parse_self_score(get_self_score_value(academic))
                if academic_self_score is None:
                    return {"success": False, "message": "学术交流缺少自评分，请补充后提交"}
                new_academic = AcademicExchange(
                    achievement_id=achievement_id,
                    name=academic.get("name"),
                    participate_type=participate_types[int(academic.get("typeIndex", 0))],
                    exchange_date=academic.get("date"),
                    self_score=academic_self_score,
                    review_status="pending"
                )
                db.add(new_academic)
                db.commit()
                db.refresh(new_academic)
                all_docs = list(academic.get("documents", []) or []) + list(academic.get("images", []) or [])
                append_documents_for_item(db, achievement_id, "academic", new_academic.id, all_docs)

        for volunteer in volunteer_items:
            if volunteer.get("project_name"):
                volunteer_self_score = parse_self_score(get_self_score_value(volunteer))
                if volunteer_self_score is None:
                    return {"success": False, "message": "志愿服务缺少自评分，请补充后提交"}
                new_volunteer = VolunteerService(
                    achievement_id=achievement_id,
                    project_name=volunteer.get("project_name"),
                    hours=int(volunteer.get("hours", 0)),
                    service_date=volunteer.get("date"),
                    self_score=volunteer_self_score,
                    review_status="pending"
                )
                db.add(new_volunteer)
                db.commit()
                db.refresh(new_volunteer)
                all_docs = list(volunteer.get("documents", []) or []) + list(volunteer.get("images", []) or [])
                append_documents_for_item(db, achievement_id, "volunteer", new_volunteer.id, all_docs)

        award_levels = ['校级', '市级', '省级', '国家级', '国际级']
        for award in award_items:
            if award.get("name"):
                award_self_score = parse_self_score(get_self_score_value(award))
                if award_self_score is None:
                    return {"success": False, "message": "获奖荣誉缺少自评分，请补充后提交"}
                new_award = Award(
                    achievement_id=achievement_id,
                    name=award.get("name"),
                    level=award_levels[int(award.get("levelIndex", 0))],
                    award_date=award.get("date"),
                    self_score=award_self_score,
                    review_status="pending"
                )
                db.add(new_award)
                db.commit()
                db.refresh(new_award)
                all_docs = list(award.get("documents", []) or []) + list(award.get("images", []) or [])
                append_documents_for_item(db, achievement_id, "award", new_award.id, all_docs)

        for custom_item in custom_items:
            try:
                type_id = int(custom_item.get("type_id"))
            except Exception:
                continue
            type_model = db.query(AchievementType).filter(AchievementType.id == type_id, AchievementType.is_active == True).first()
            if not type_model:
                continue
            content_obj = custom_item.get("content", {})
            if not isinstance(content_obj, dict):
                continue
            custom_self_score = parse_self_score(get_self_score_value(custom_item))
            if custom_self_score is None:
                return {"success": False, "message": "自定义成果缺少自评分，请补充后提交"}
            new_custom = CustomAchievement(
                achievement_id=achievement_id,
                type_id=type_id,
                content_json=json.dumps(content_obj, ensure_ascii=False),
                self_score=custom_self_score,
                review_status="pending"
            )
            db.add(new_custom)
            db.commit()
            db.refresh(new_custom)
            all_docs = list(custom_item.get("documents", []) or [])
            append_documents_for_item(db, achievement_id, "custom", new_custom.id, all_docs)

        db.commit()
        return {
            "success": True,
            "message": "成果提交成功，等待审核",
            "achievement_id": achievement_id,
            "student_id": student_id
        }
    except Exception as e:
        db.rollback()
        print(f"提交成果失败：{str(e)}")
        return {
            "success": False,
            "message": f"提交失败：{str(e)}"
        }



# 6. 管理端：获取成果列表（支持学号/审核状态筛选）
@app.get("/admin/achievements")
async def get_achievements(
    page: int = 1,
    size: int = 10,
    audit_status: Optional[bool] = None,
    student_id: Optional[str] = None,
    db: Session = Depends(get_db)
):
    try:
        # 构建查询条件
        query = db.query(StudentAchievement)
        
        # 审核状态筛选
        if audit_status is not None:
            query = query.filter(StudentAchievement.audit_status == audit_status)
        
        # 学号筛选
        if student_id and student_id.strip():
            query = query.filter(StudentAchievement.student_id == student_id.strip())
        
        # 分页处理
        total = query.count()
        achievements = query.order_by(StudentAchievement.create_time.desc()).offset((page-1)*size).limit(size).all()
        
        # 组装返回数据
        result = []
        for item in achievements:
            result.append({
                "id": item.id,
                "student_id": item.student_id,
                "openid": item.openid,
                "create_time": item.create_time.strftime("%Y-%m-%d %H:%M:%S") if item.create_time else "",
                "audit_status": item.audit_status,
                "audit_note": item.audit_note or "",
                "audit_time": item.audit_time.strftime("%Y-%m-%d %H:%M:%S") if item.audit_time else "",
                "overall_score": item.overall_score,
                "review_completed": bool(item.review_completed)
            })
        
        return {
            "code": 200,
            "data": {
                "list": result,
                "total": total,
                "page": page,
                "size": size
            },
            "message": "查询成功"
        }
    except Exception as e:
        print(f"查询成果列表失败：{str(e)}")
        return {
            "code": 500,
            "data": {
                "list": [],
                "total": 0,
                "page": page,
                "size": size
            },
            "message": f"查询失败：{str(e)}"
        }

# 7. 管理端：获取成果详情
@app.get("/admin/achievements/{achievement_id}")
async def get_achievement_detail(
    achievement_id: int,
    db: Session = Depends(get_db)
):
    try:
        achievement = db.query(StudentAchievement).filter(StudentAchievement.id == achievement_id).first()
        if not achievement:
            return {
                "code": 404,
                "data": None,
                "message": "成果记录不存在"
            }
        papers = db.query(Paper).filter(Paper.achievement_id == achievement_id).all()
        policies = db.query(PolicyReport).filter(PolicyReport.achievement_id == achievement_id).all()
        academics = db.query(AcademicExchange).filter(AcademicExchange.achievement_id == achievement_id).all()
        volunteers = db.query(VolunteerService).filter(VolunteerService.achievement_id == achievement_id).all()
        awards = db.query(Award).filter(Award.achievement_id == achievement_id).all()
        customs = db.query(CustomAchievement).filter(CustomAchievement.achievement_id == achievement_id).all()

        def format_items(items, item_type):
            result = []
            for item in items:
                item_data = {
                    "id": item.id,
                    "documents": get_documents_for_item(db, item_type, item.id),
                    "self_score": item.self_score,
                    "review_score": item.review_score,
                    "rescore_score": item.rescore_score,
                    "review_status": item.review_status or "pending",
                    "review_comment": item.review_comment or "",
                    "student_agree": item.student_agree,
                    "student_feedback_comment": item.student_feedback_comment or "",
                    "rescore_comment": item.rescore_comment or ""
                }
                if item_type == "paper":
                    item_data.update({
                        "title": item.title,
                        "journal": item.journal,
                        "publish_date": item.publish_date
                    })
                elif item_type == "policy":
                    item_data.update({
                        "title": item.title,
                        "adopt_unit": item.adopt_unit,
                        "submit_date": item.submit_date
                    })
                elif item_type == "academic":
                    item_data.update({
                        "name": item.name,
                        "participate_type": item.participate_type,
                        "exchange_date": item.exchange_date
                    })
                elif item_type == "volunteer":
                    item_data.update({
                        "project_name": item.project_name,
                        "hours": item.hours,
                        "service_date": item.service_date
                    })
                elif item_type == "award":
                    item_data.update({
                        "name": item.name,
                        "level": item.level,
                        "award_date": item.award_date
                    })
                elif item_type == "custom":
                    type_info = db.query(AchievementType).filter(AchievementType.id == item.type_id).first()
                    item_data.update({
                        "type_id": item.type_id,
                        "type_name": type_info.name if type_info else f"类型{item.type_id}",
                        "content": json.loads(item.content_json or "{}")
                    })
                result.append(item_data)
            return result
        paper_data = format_items(papers, "paper")
        policy_data = format_items(policies, "policy")
        academic_data = format_items(academics, "academic")
        volunteer_data = format_items(volunteers, "volunteer")
        award_data = format_items(awards, "award")
        custom_data = format_items(customs, "custom")
        type_summaries = []
        paper_summary = build_type_summary("论文成果", "paper", papers)
        policy_summary = build_type_summary("资政报告", "policy", policies)
        academic_summary = build_type_summary("学术交流", "academic", academics)
        volunteer_summary = build_type_summary("志愿服务", "volunteer", volunteers)
        award_summary = build_type_summary("获奖荣誉", "award", awards)
        custom_summary = build_type_summary("自定义成果", "custom", customs)
        for summary in [paper_summary, policy_summary, academic_summary, volunteer_summary, award_summary, custom_summary]:
            if summary:
                docs = []
                source = {
                    "paper": paper_data,
                    "policy": policy_data,
                    "academic": academic_data,
                    "volunteer": volunteer_data,
                    "award": award_data,
                    "custom": custom_data
                }.get(summary["type"], [])
                for item in source:
                    docs.extend(item.get("documents", []))
                summary["documents"] = docs
                type_summaries.append(summary)
        detail_data = {
            "id": achievement.id,
            "student_id": achievement.student_id,
            "create_time": achievement.create_time.strftime("%Y-%m-%d %H:%M:%S") if achievement.create_time else "",
            "audit_status": achievement.audit_status,
            "audit_note": achievement.audit_note or "",
            "audit_time": achievement.audit_time.strftime("%Y-%m-%d %H:%M:%S") if achievement.audit_time else "",
            "overall_score": achievement.overall_score,
            "review_completed": bool(achievement.review_completed),
            "papers": paper_data,
            "policies": policy_data,
            "academics": academic_data,
            "volunteers": volunteer_data,
            "awards": award_data,
            "customs": custom_data,
            "type_summaries": type_summaries
        }
        achievement.audit_status = calculate_review_completed(detail_data)
        if achievement.audit_status and not achievement.audit_time:
            achievement.audit_time = datetime.now()
        db.commit()
        return {
            "code": 200,
            "data": detail_data,
            "message": "查询成功"
        }
    except Exception as e:
        print(f"查询成果详情失败：{str(e)}")
        return {
            "code": 500,
            "data": None,
            "message": f"查询失败：{str(e)}"
        }

@app.post("/admin/review/{item_type}/{item_id}")
async def review_single_item(
    item_type: str,
    item_id: int,
    data: dict = Body(...),
    db: Session = Depends(get_db)
):
    item_type = (item_type or "").lower()
    model_info = REVIEW_MODEL_MAP.get(item_type)
    if not model_info:
        return {"code": 400, "message": "不支持的成果类型", "data": None}
    model_cls = model_info[0]
    item = db.query(model_cls).filter(model_cls.id == item_id).first()
    if not item:
        return {"code": 404, "message": "成果项不存在", "data": None}
    score = data.get("score")
    if score is None:
        return {"code": 400, "message": "请填写评分", "data": None}
    try:
        numeric_score = float(score)
    except Exception:
        return {"code": 400, "message": "评分格式错误", "data": None}
    if numeric_score < 0 or numeric_score > 100:
        return {"code": 400, "message": "评分范围必须在0到100", "data": None}
    item.review_score = numeric_score
    item.rescore_score = None
    item.review_comment = str(data.get("review_comment") or "")
    item.rescore_comment = str(data.get("rescore_comment") or "")
    item.review_status = "reviewed"
    item.review_time = datetime.now()
    achievement = db.query(StudentAchievement).filter(StudentAchievement.id == item.achievement_id).first()
    if achievement:
        recalculate_achievement_score(db, achievement)
        achievement.audit_status = bool(achievement.review_completed)
        if achievement.audit_status:
            achievement.audit_time = datetime.now()
    db.commit()
    return {
        "code": 200,
        "message": "评分成功",
        "data": {
            "item_type": item_type,
            "item_id": item_id,
            "score": item.review_score,
            "rescore_score": item.rescore_score,
            "overall_score": achievement.overall_score if achievement else None,
            "review_completed": bool(achievement.review_completed) if achievement else False
        }
    }

@app.post("/admin/review-type/{achievement_id}/{item_type}")
async def review_type_score(
    achievement_id: int,
    item_type: str,
    data: dict = Body(...),
    db: Session = Depends(get_db)
):
    item_type = (item_type or "").lower()
    items = get_type_items(db, achievement_id, item_type)
    if not items:
        return {"code": 404, "data": None, "message": "该类型成果不存在"}
    score = data.get("score")
    rescore_score = data.get("rescore_score")
    if score is None and rescore_score is None:
        return {"code": 400, "data": None, "message": "请填写评分"}
    try:
        score = float(score) if score is not None else None
    except Exception:
        return {"code": 400, "data": None, "message": "评分格式错误"}
    try:
        rescore_score = float(rescore_score) if rescore_score is not None else None
    except Exception:
        return {"code": 400, "data": None, "message": "复核评分格式错误"}
    if score is not None and (score < 0 or score > 100):
        return {"code": 400, "data": None, "message": "评分范围必须在0到100"}
    if rescore_score is not None and (rescore_score < 0 or rescore_score > 100):
        return {"code": 400, "data": None, "message": "复核评分范围必须在0到100"}
    review_comment = str(data.get("review_comment") or "")
    rescore_comment = str(data.get("rescore_comment") or "")
    was_disagreed = any((item.review_status or "") == "disagreed" for item in items)
    has_rescore_input = bool(rescore_comment.strip()) or (rescore_score is not None)
    should_rescore = bool(was_disagreed or has_rescore_input)
    final_review_score = score
    if final_review_score is None:
        existed_review_score = items[0].review_score
        final_review_score = float(existed_review_score) if existed_review_score is not None else None
    final_rescore_score = None
    if should_rescore:
        final_rescore_score = rescore_score if rescore_score is not None else final_review_score
        if final_rescore_score is None:
            return {"code": 400, "data": None, "message": "请填写复核评分"}
    elif final_review_score is None:
        return {"code": 400, "data": None, "message": "请填写评分"}
    for item in items:
        if final_review_score is not None:
            item.review_score = final_review_score
        item.rescore_score = final_rescore_score
        item.review_comment = review_comment
        item.rescore_comment = rescore_comment
        item.review_time = datetime.now()
        item.review_status = "rescored" if should_rescore else "reviewed"
    achievement = db.query(StudentAchievement).filter(StudentAchievement.id == achievement_id).first()
    if achievement:
        recalculate_achievement_score(db, achievement)
        achievement.audit_status = bool(achievement.review_completed)
        if achievement.audit_status:
            achievement.audit_time = datetime.now()
    db.commit()
    return {
        "code": 200,
        "data": {
            "achievement_id": achievement_id,
            "item_type": item_type,
            "score": final_review_score,
            "rescore_score": final_rescore_score,
            "review_status": "rescored" if should_rescore else "reviewed",
            "overall_score": achievement.overall_score if achievement else None
        },
        "message": "类型评分成功"
    }

@app.post("/admin/review-item/{item_type}/{item_id}")
async def review_item_score(
    item_type: str,
    item_id: int,
    data: dict = Body(...),
    db: Session = Depends(get_db)
):
    item_type = (item_type or "").lower()
    model_info = REVIEW_MODEL_MAP.get(item_type)
    if not model_info:
        return {"code": 400, "data": None, "message": "成果类型错误"}
    model_cls = model_info[0]
    
    item = db.query(model_cls).filter(model_cls.id == item_id).first()
    if not item:
        return {"code": 404, "data": None, "message": "成果项不存在"}

    score = data.get("score")
    rescore_score = data.get("rescore_score")
    review_comment = str(data.get("review_comment") or "")
    rescore_comment = str(data.get("rescore_comment") or "")

    try:
        if score is not None:
            score = float(score)
            if score < 0 or score > 100: return {"code": 400, "data": None, "message": "评分范围必须在0到100"}
            item.review_score = score
        if rescore_score is not None:
            rescore_score = float(rescore_score)
            if rescore_score < 0 or rescore_score > 100: return {"code": 400, "data": None, "message": "复核评分范围必须在0到100"}
            item.rescore_score = rescore_score
    except Exception:
        return {"code": 400, "data": None, "message": "评分格式错误"}

    item.review_comment = review_comment
    item.rescore_comment = rescore_comment
    item.review_time = datetime.now()
    
    if item.rescore_score is not None:
        item.review_status = "rescored"
    elif item.review_score is not None:
        item.review_status = "reviewed"
    
    achievement = db.query(StudentAchievement).filter(StudentAchievement.id == item.achievement_id).first()
    if achievement:
        recalculate_achievement_score(db, achievement)

    db.commit()
    return {"code": 200, "data": {"item_id": item_id, "score": item.review_score, "rescore_score": item.rescore_score}, "message": "评分成功"}

@app.get("/admin/score-formula")
async def get_score_formula(db: Session = Depends(get_db)):
    formula = get_or_init_score_formula(db)
    return {
        "code": 200,
        "data": {
            "id": formula.id,
            "weights": parse_weights(formula.weights_json),
            "update_time": formula.update_time.strftime("%Y-%m-%d %H:%M:%S") if formula.update_time else ""
        },
        "message": "查询成功"
    }

@app.put("/admin/score-formula")
async def update_score_formula(data: dict = Body(...), db: Session = Depends(get_db)):
    weights = data.get("weights", {})
    if not isinstance(weights, dict):
        return {"code": 400, "data": None, "message": "权重格式错误"}
    normalized = {}
    for key in ["paper", "policy", "academic", "volunteer", "award", "custom"]:
        try:
            normalized[key] = float(weights.get(key, 0))
        except Exception:
            normalized[key] = 0.0
    formula = get_or_init_score_formula(db)
    formula.weights_json = json.dumps(normalized, ensure_ascii=False)
    formula.update_time = datetime.now()
    db.commit()
    for item in db.query(StudentAchievement).all():
        recalculate_achievement_score(db, item)
        item.audit_status = bool(item.review_completed)
    db.commit()
    return {"code": 200, "data": {"weights": normalized}, "message": "更新成功"}

@app.get("/admin/achievement-types")
async def get_achievement_types(db: Session = Depends(get_db)):
    items = db.query(AchievementType).order_by(AchievementType.id.asc()).all()
    return {
        "code": 200,
        "data": {
            "list": [
                {
                    "id": item.id,
                    "name": item.name,
                    "fields": json.loads(item.fields_json or "[]"),
                    "is_active": bool(item.is_active)
                }
                for item in items
            ]
        },
        "message": "查询成功"
    }

@app.post("/admin/achievement-types")
async def create_achievement_type(data: dict = Body(...), db: Session = Depends(get_db)):
    name = str(data.get("name") or "").strip()
    fields = data.get("fields", [])
    if not name:
        return {"code": 400, "data": None, "message": "类型名称不能为空"}
    if not isinstance(fields, list):
        return {"code": 400, "data": None, "message": "字段配置格式错误"}
    exists = db.query(AchievementType).filter(AchievementType.name == name).first()
    if exists:
        return {"code": 400, "data": None, "message": "类型名称已存在"}
    item = AchievementType(
        name=name,
        fields_json=json.dumps(fields, ensure_ascii=False),
        is_active=bool(data.get("is_active", True)),
        create_time=datetime.now(),
        update_time=datetime.now()
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return {"code": 200, "data": {"id": item.id}, "message": "新增成功"}

@app.put("/admin/achievement-types/{type_id}")
async def update_achievement_type(type_id: int, data: dict = Body(...), db: Session = Depends(get_db)):
    item = db.query(AchievementType).filter(AchievementType.id == type_id).first()
    if not item:
        return {"code": 404, "data": None, "message": "类型不存在"}
    if "name" in data and str(data.get("name") or "").strip():
        item.name = str(data.get("name")).strip()
    if "fields" in data and isinstance(data.get("fields"), list):
        item.fields_json = json.dumps(data.get("fields"), ensure_ascii=False)
    if "is_active" in data:
        item.is_active = bool(data.get("is_active"))
    item.update_time = datetime.now()
    db.commit()
    return {"code": 200, "data": {"id": item.id}, "message": "更新成功"}

@app.post("/student/achievements/{achievement_id}/feedback")
async def submit_student_feedback(
    achievement_id: int,
    item_type: str = Body(...),
    item_id: int = Body(...),
    agree: bool = Body(...),
    comment: str = Body(""),
    current_student: StudentUser = Depends(get_current_student),
    db: Session = Depends(get_db)
):
    achievement = db.query(StudentAchievement).filter(
        StudentAchievement.id == achievement_id,
        StudentAchievement.student_id == current_student.student_id
    ).first()
    if not achievement:
        return {"code": 404, "data": None, "message": "成果记录不存在"}
    model_info = REVIEW_MODEL_MAP.get((item_type or "").lower())
    if not model_info:
        return {"code": 400, "data": None, "message": "成果类型错误"}
    model_cls = model_info[0]
    target_items = db.query(model_cls).filter(
        model_cls.achievement_id == achievement_id
    ).all()
    if not target_items:
        return {"code": 404, "data": None, "message": "成果项不存在"}
    if any(item.student_agree is not None for item in target_items):
        return {"code": 400, "data": None, "message": "该类型成果已反馈，不可重复提交"}
    selected_item = None
    for item in target_items:
        if item.id == item_id:
            selected_item = item
            break
    if not selected_item:
        return {"code": 404, "data": None, "message": "成果项不存在"}
    is_agree = bool(agree)
    feedback_comment = str(comment or "")
    for item in target_items:
        item.student_agree = is_agree
        item.student_feedback_comment = feedback_comment
        item.feedback_time = datetime.now()
        if is_agree:
            item.review_status = "agreed"
        elif (item.review_status or "") in ["reviewed", "agreed"]:
            item.review_status = "disagreed"
    recalculate_achievement_score(db, achievement)
    achievement.audit_status = bool(achievement.review_completed)
    db.commit()
    return {"code": 200, "data": {"item_id": item_id, "agree": is_agree}, "message": "反馈成功"}

@app.get("/student/achievements")
async def get_student_achievements(
    current_student: StudentUser = Depends(get_current_student),
    db: Session = Depends(get_db)
):
    achievements = db.query(StudentAchievement).filter(
        StudentAchievement.student_id == current_student.student_id
    ).order_by(StudentAchievement.create_time.desc()).all()
    response_list = []
    for achievement in achievements:
        detail_res = await get_achievement_detail(achievement.id, db)
        if detail_res.get("code") == 200:
            response_list.append(detail_res.get("data"))
    return {"code": 200, "data": {"list": response_list}, "message": "查询成功"}

from dotenv import load_dotenv

load_dotenv()

@app.post("/student/ai-extract-fields")
async def ai_extract_fields(
    achievement_type: str = Body(...),
    document_paths: List[str] = Body(...),
    db: Session = Depends(get_db)
):
    # 1. 获取API Key
    api_key = os.getenv("BAILIAN_API_KEY", "").strip()
    if not api_key:
        return {"code": 200, "data": {"suggestions": {}, "enabled": False}, "message": "AI提取未启用"}
    
    dashscope.api_key = api_key

    # 2. 定义字段映射
    field_maps = {
        "论文成果": [
            {"key": "title", "label": "论文标题"},
            {"key": "journal", "label": "发表期刊"},
            {"key": "date", "label": "发表日期(YYYY-MM-DD)"}
        ],
        "paper": [
            {"key": "title", "label": "论文标题"},
            {"key": "journal", "label": "发表期刊"},
            {"key": "date", "label": "发表日期(YYYY-MM-DD)"}
        ],
        "资政报告": [
            {"key": "title", "label": "报告标题"},
            {"key": "adopt_unit", "label": "采纳单位"},
            {"key": "date", "label": "提交日期(YYYY-MM-DD)"}
        ],
        "policy": [
            {"key": "title", "label": "报告标题"},
            {"key": "adopt_unit", "label": "采纳单位"},
            {"key": "date", "label": "提交日期(YYYY-MM-DD)"}
        ],
        "学术交流": [
            {"key": "name", "label": "交流名称"},
            {"key": "typeIndex", "label": "参与类型(学术会议/学术讲座/其他)"},
            {"key": "date", "label": "交流日期(YYYY-MM-DD)"}
        ],
        "academic": [
            {"key": "name", "label": "交流名称"},
            {"key": "typeIndex", "label": "参与类型(学术会议/学术讲座/其他)"},
            {"key": "date", "label": "交流日期(YYYY-MM-DD)"}
        ],
        "志愿服务": [
            {"key": "project_name", "label": "项目名称"},
            {"key": "hours", "label": "服务时长(数字)"},
            {"key": "date", "label": "服务日期(YYYY-MM-DD)"}
        ],
        "volunteer": [
            {"key": "project_name", "label": "项目名称"},
            {"key": "hours", "label": "服务时长(数字)"},
            {"key": "date", "label": "服务日期(YYYY-MM-DD)"}
        ],
        "获奖荣誉": [
            {"key": "name", "label": "奖项名称"},
            {"key": "levelIndex", "label": "奖项级别(国家级/省部级/校级/院级)"},
            {"key": "date", "label": "获奖日期(YYYY-MM-DD)"}
        ],
        "award": [
            {"key": "name", "label": "奖项名称"},
            {"key": "levelIndex", "label": "奖项级别(国家级/省部级/校级/院级)"},
            {"key": "date", "label": "获奖日期(YYYY-MM-DD)"}
        ]
    }

    target_fields = field_maps.get(achievement_type)
    if not target_fields:
        return {"code": 200, "data": {"suggestions": {}, "enabled": True}, "message": f"未知的成果类型: {achievement_type}"}

    # 3. 筛选文件
    image_extensions = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".gif"}
    doc_extensions = {".pdf", ".docx", ".txt"}
    
    image_files = []
    doc_files = []
    
    server_base_url = os.getenv("SERVER_BASE_URL", "https://api.aipro.ren").rstrip("/")
    
    for path in (document_paths or []):
        if not path:
            continue
        file_ext = os.path.splitext(path)[1].lower()
        
        # 优先使用本地路径
        local_path = os.path.join(UPLOAD_DIR, path)
        if not os.path.exists(local_path):
            continue
        abs_path = os.path.abspath(local_path)
        
        if file_ext in image_extensions:
            image_files.append(f"file://{abs_path}")
        elif file_ext in doc_extensions:
            doc_files.append(abs_path)

    if not image_files and not doc_files:
        return {"code": 200, "data": {"suggestions": {}, "enabled": True}, "message": "未找到支持的文件(图片/PDF/Word/TXT)"}

    # 4. 提取文档内容
    doc_content = ""
    for f in doc_files:
        ext = os.path.splitext(f)[1].lower()
        try:
            text = ""
            if ext == ".pdf":
                reader = PdfReader(f)
                for page in reader.pages:
                    text += (page.extract_text() or "") + "\n"
            elif ext == ".docx":
                doc = Document(f)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            elif ext == ".txt":
                with open(f, "r", encoding="utf-8") as tf:
                    text = tf.read()
            
            if text.strip():
                doc_content += f"\n--- 文档内容 ({os.path.basename(f)}) ---\n{text[:50000]}\n" # Limit 50k chars per doc
        except Exception as e:
            print(f"Error reading {f}: {e}")

    # 5. 构建Prompt
    fields_desc = ", ".join([f"{f['key']}({f['label']})" for f in target_fields])
    prompt = f"请提取以下字段信息：{fields_desc}。请直接返回JSON格式数据，不要包含Markdown标记。"
    
    if doc_content:
        prompt += f"\n\n参考文档内容:\n{doc_content}"

    # 6. 调用API
    try:
        if image_files:
            # 使用多模态模型 (qwen-vl-plus)
            messages = [{"role": "user", "content": []}]
            for img_url in image_files:
                messages[0]["content"].append({"image": img_url})
            messages[0]["content"].append({"text": prompt})
            
            model_name = "qwen-vl-plus"
            response = MultiModalConversation.call(model=model_name, messages=messages)
            
            # Response parsing for VL model
            if response.status_code == 200:
                result_text = ""
                if response.output and response.output.choices:
                    message_content = response.output.choices[0].message.content
                    if isinstance(message_content, list):
                        for item in message_content:
                            if "text" in item:
                                result_text += item["text"]
                    elif isinstance(message_content, str):
                        result_text = message_content
                content = result_text
            else:
                return {"code": 500, "data": {"suggestions": {}, "enabled": True}, "message": f"VL模型调用失败: {response.code} - {response.message}"}
                
        else:
            # 纯文本模型 (qwen-plus)
            messages = [{"role": "user", "content": prompt}]
            model_name = "qwen-plus" 
            
            response = Generation.call(model=model_name, messages=messages, result_format='message')
            
            # Response parsing for Generation model
            if response.status_code == 200:
                content = response.output.choices[0].message.content
            else:
                return {"code": 500, "data": {"suggestions": {}, "enabled": True}, "message": f"LLM调用失败: {response.code} - {response.message}"}

        # Common JSON parsing
        if content.startswith("```json"):
            content = content[7:]
        if content.endswith("```"):
            content = content[:-3]
        
        suggestions = json.loads(content.strip())
        return {"code": 200, "data": {"suggestions": suggestions, "enabled": True}, "message": "提取成功"}

    except Exception as e:
        return {"code": 500, "data": {"suggestions": {}, "enabled": True}, "message": f"提取失败: {str(e)}"}
# 9. 静态文件访问（图片预览）
@app.get("/uploads/{file_name}")
async def get_uploaded_file(file_name: str):
    file_path = os.path.join(UPLOAD_DIR, file_name)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")
    return FileResponse(file_path)

# ========== 启动时创建数据库表 ==========
@app.on_event("startup")
async def startup():
    Base.metadata.create_all(bind=engine)
    ensure_student_user_schema()
    ensure_achievement_schema()
    db = SessionLocal()
    try:
        if db.query(AdminPermission).count() == 0:
            default_permissions = [
                {"name": "用户管理", "key": "user:manage", "description": "管理系统用户"},
                {"name": "角色管理", "key": "role:manage", "description": "管理系统角色"},
                {"name": "权限管理", "key": "permission:manage", "description": "管理系统权限"},
                {"name": "学生管理", "key": "student:manage", "description": "管理学生信息"}
            ]
            db.add_all([
                AdminPermission(
                    name=item["name"],
                    key=item["key"],
                    description=item["description"],
                    create_time=datetime.now()
                )
                for item in default_permissions
            ])
            db.commit()
        if db.query(AdminRole).count() == 0:
            permission_keys = [item.key for item in db.query(AdminPermission).all()]
            roles = [
                AdminRole(
                    name="管理员",
                    description="系统管理员，拥有所有权限",
                    permissions=json.dumps(permission_keys, ensure_ascii=False),
                    create_time=datetime.now()
                ),
                AdminRole(
                    name="审核员",
                    description="负责审核学生成果",
                    permissions=json.dumps(["student:manage"], ensure_ascii=False),
                    create_time=datetime.now()
                ),
                AdminRole(
                    name="普通用户",
                    description="普通系统用户",
                    permissions=json.dumps([], ensure_ascii=False),
                    create_time=datetime.now()
                )
            ]
            db.add_all(roles)
            db.commit()
        get_or_init_score_formula(db)
        admin_user = db.query(AdminUser).filter(AdminUser.username == "admin").first()
        if not admin_user:
            admin_role = db.query(AdminRole).filter(AdminRole.name == "管理员").first()
            password_bytes = "123456".encode("utf-8")
            hashed_password = bcrypt.hashpw(password_bytes, bcrypt.gensalt()).decode("utf-8")
            admin_user = AdminUser(
                username="admin",
                name="管理员",
                email="admin@example.com",
                password=hashed_password,
                role_id=admin_role.id if admin_role else None,
                is_active=True,
                create_time=datetime.now()
            )
            db.add(admin_user)
            db.commit()
        whitelist_student = db.query(StudentUser).filter(StudentUser.student_id == "20260001").first()
        if not whitelist_student:
            default_password = "123456"
            password_bytes = default_password.encode("utf-8")
            hashed_password = bcrypt.hashpw(password_bytes, bcrypt.gensalt()).decode("utf-8")
            whitelist_student = StudentUser(
                name="测试学生",
                student_id="20260001",
                password=hashed_password,
                default_password=default_password,
                create_time=datetime.now(),
                update_time=datetime.now(),
                is_active=True,
                is_whitelisted=True,
                must_change_password=True
            )
            db.add(whitelist_student)
            db.commit()
        if db.query(AchievementType).count() == 0:
            default_types = [
                {"name": "论文成果", "fields": [{"key": "title", "label": "论文标题"}, {"key": "journal", "label": "发表期刊"}, {"key": "date", "label": "发表日期"}]},
                {"name": "资政报告", "fields": [{"key": "title", "label": "报告标题"}, {"key": "adopt_unit", "label": "采纳单位"}, {"key": "date", "label": "提交日期"}]},
                {"name": "学术交流", "fields": [{"key": "name", "label": "交流名称"}, {"key": "typeIndex", "label": "参与类型"}, {"key": "date", "label": "交流日期"}]},
                {"name": "志愿服务", "fields": [{"key": "project_name", "label": "项目名称"}, {"key": "hours", "label": "服务时长"}, {"key": "date", "label": "服务日期"}]},
                {"name": "获奖荣誉", "fields": [{"key": "name", "label": "奖项名称"}, {"key": "levelIndex", "label": "奖项级别"}, {"key": "date", "label": "获奖日期"}]}
            ]
            for entry in default_types:
                db.add(AchievementType(
                    name=entry["name"],
                    fields_json=json.dumps(entry["fields"], ensure_ascii=False),
                    is_active=True,
                    create_time=datetime.now(),
                    update_time=datetime.now()
                ))
            db.commit()
    finally:
        db.close()

# ========== 新增：管理端-获取提交成果的学生列表 ==========
@app.get("/admin/students", response_model=dict)
async def get_student_list(
    page: int = 1,
    size: int = 10,
    student_id: str = None,
    name: str = None,
    audit_status: str = None,
    db: Session = Depends(get_db)
):
    try:
        # 1. 查询所有提交过成果的学生（去重）
        sub_query = db.query(StudentAchievement.student_id).distinct().subquery()
        query = db.query(StudentUser)
        query = query.filter(StudentUser.student_id.in_(sub_query))
        
        # 2. 添加筛选条件
        if student_id and student_id.strip():
            query = query.filter(StudentUser.student_id.like(f"%{student_id.strip()}%"))
        if name and name.strip():
            query = query.filter(StudentUser.name.like(f"%{name.strip()}%"))
        
        # 3. 查询候选学生（后续按状态筛选再分页）
        students = query.all()
        
        # 4. 组装数据（补充成果数和审核状态）
        result = []
        for student in students:
            # 统计该学生的成果数
            achievement_count = db.query(StudentAchievement).filter(
                StudentAchievement.student_id == student.student_id
            ).count()
            
            # 最后提交时间
            last_achievement = db.query(StudentAchievement).filter(
                StudentAchievement.student_id == student.student_id
            ).order_by(StudentAchievement.create_time.desc()).first()
            last_submit_time = last_achievement.create_time.strftime("%Y-%m-%d %H:%M:%S") if last_achievement else ""
            latest_overall_score = last_achievement.overall_score if last_achievement else None
            
            lifecycle_status = calculate_achievement_lifecycle_status(db, last_achievement) if last_achievement else "已提交"
            
            result.append({
                "student_id": student.student_id,
                "name": student.name,
                "submit_count": achievement_count,
                "last_submit_time": last_submit_time,
                "audit_status": lifecycle_status,
                "latest_overall_score": latest_overall_score
            })
        if audit_status and str(audit_status).strip():
            result = [item for item in result if item["audit_status"] == str(audit_status).strip()]
        total = len(result)
        page_result = result[(page - 1) * size: page * size]

        return {
            "code": 200,
            "data": {
                "list": page_result,
                "total": total,
                "page": page,
                "size": size
            },
            "message": "查询成功"
        }
    except Exception as e:
        print(f"查询学生列表失败：{str(e)}")
        return {
            "code": 500,
            "data": {
                "list": [],
                "total": 0
            },
            "message": f"查询失败：{str(e)}"
        }

# ========== 新增：管理端-获取学生基础信息 ==========
@app.get("/admin/students/{student_id}", response_model=dict)
async def get_student_info(
    student_id: str,
    db: Session = Depends(get_db)
):
    try:
        # 查询学生基础信息
        student = db.query(StudentUser).filter(StudentUser.student_id == student_id).first()
        if not student:
            return {
                "code": 404,
                "data": None,
                "message": "学生不存在"
            }
        
        # 统计该学生的成果总数
        total_achievements = db.query(StudentAchievement).filter(
            StudentAchievement.student_id == student_id
        ).count()
        
        # 整体审核状态（是否全部审核）
        all_audited = db.query(StudentAchievement).filter(
            StudentAchievement.student_id == student_id,
            StudentAchievement.audit_status == False
        ).count() == 0
        
        return {
            "code": 200,
            "data": {
                "student_id": student.student_id,
                "name": student.name,
                "total_achievements": total_achievements,
                "audit_status": all_audited,
                "latest_overall_score": db.query(StudentAchievement).filter(
                    StudentAchievement.student_id == student_id
                ).order_by(StudentAchievement.create_time.desc()).first().overall_score if total_achievements > 0 else None
            },
            "message": "查询成功"
        }
    except Exception as e:
        print(f"查询学生信息失败：{str(e)}")
        return {
            "code": 500,
            "data": None,
            "message": f"查询失败：{str(e)}"
        }


# ========== 主函数（直接运行） ==========
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)
